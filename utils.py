import re
import pandas as pd
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
import ollama

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != '']
    return paragraphs

def detect_ai_markers(text):
    ai_markers = [
        r"\b(as an AI language model|as an AI model|I don't have personal opinions)\b",
        r"\b(large language model|LLM|GPT-4|ChatGPT)\b",
        r"\b(in conclusion|overall|in summary|it is important to note)\b",
        r"\b(one might argue|some might suggest|many people believe)\b",
        r"\b(however|therefore|furthermore|additionally|on top of that)\b"
    ]
    for marker in ai_markers:
        if re.search(marker, text, re.I):
            return True
    return False

def analyze_text_with_llm(text):
    model = ollama.load_model('gpt-4')  # Load the GPT-4 model
    response = model.analyze(text)
    return response

def process_document(file_path):
    paragraphs = extract_text_from_docx(file_path)
    df = pd.DataFrame(paragraphs, columns=['Paragraph'])
    df['AI_Marker_Detected'] = df['Paragraph'].apply(detect_ai_markers)

    vectorizer = CountVectorizer().fit_transform(df['Paragraph'])
    vectors = vectorizer.toarray()
    cosine_matrix = cosine_similarity(vectors)

    plagiarism_threshold = 0.75
    plagiarism_flags = []
    for i in range(len(cosine_matrix)):
        similarities = cosine_matrix[i]
        high_similarity = [j for j in range(len(similarities)) if similarities[j] > plagiarism_threshold and i != j]
        plagiarism_flags.append(bool(high_similarity))

    df['Possible_Plagiarism'] = plagiarism_flags
    return df, cosine_matrix

def generate_report(df, cosine_matrix):
    ai_generated_percentage = np.mean(df['AI_Marker_Detected']) * 100
    plagiarism_percentage = np.mean(df['Possible_Plagiarism']) * 100

    report = f"""
    Analysis Report:
    ----------------
    AI-Generated Markers: {ai_generated_percentage:.2f}%
    Possible Plagiarism: {plagiarism_percentage:.2f}%

    Suggestions to reduce AI writing presence:
    - Modify the highlighted keywords and strings.
    - Adjust the tonality to be more human-like.
    """
    return report

def reconstruct_document(file_path, df):
    replacement_dict = {
        "in conclusion": "to sum up",
        "overall": "generally speaking",
        "in summary": "briefly stated",
        "it is important to note": "notably",
        "on top of that": "besides",
        "furthermore": "also",
        "however": "yet",
        "therefore": "thus",
        "additionally": "moreover"
    }

    def replace_ai_phrases(match):
        return replacement_dict[match.group().lower()]

    rewritten_paragraphs = []
    for index, row in df.iterrows():
        paragraph = row['Paragraph']
        issue = row['Issue']
        
        if issue == 'AI-Generated Marker':
            paragraph = re.sub(
                r"\b(in conclusion|overall|in summary|it is important to note|on top of that|furthermore|however|therefore|additionally)\b",
                replace_ai_phrases, paragraph, flags=re.I
            )
        
        if issue == 'Possible_Plagiarism':
            sentences = re.split(r'(?<=[.!?]) +', paragraph)
            if len(sentences) > 1:
                paragraph = ' '.join(sentences[::-1])
            paragraph = re.sub(r'\b(chances|probability|likelihood)\b', 'odds', paragraph, flags=re.I)
            paragraph = re.sub(
                r'\b(design|fine-tuning|universe)\b',
                lambda x: {'design': 'structure', 'fine-tuning': 'precise arrangement', 'universe': 'cosmos'}[x.group().lower()],
                paragraph, flags=re.I
            )
        
        rewritten_paragraphs.append(paragraph)

    new_doc = Document()
    for paragraph in rewritten_paragraphs:
        new_doc.add_paragraph(paragraph)

    new_doc_path = 'path/to/your/rewritten_essay.docx'
    new_doc.save(new_doc_path)
    return new_doc_path
